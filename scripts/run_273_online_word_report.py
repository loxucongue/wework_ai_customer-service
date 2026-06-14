from __future__ import annotations

import json
import os
import sys
import time
import urllib.error
import urllib.request
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from statistics import mean
from typing import Any

import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "ai_paths"))
DEFAULT_XLSX = Path(r"C:\Users\24159\Desktop\贝颜数据\ai客服场景问题+业务逻辑.xlsx")
API_URL = os.getenv("AI_PATHS_ONLINE_REPLY_URL", "http://47.252.81.104/api/ai/reply")
INPROCESS = os.getenv("AI_PATHS_273_INPROCESS", "0") == "1"
INPUT_XLSX = Path(os.getenv("AI_PATHS_273_XLSX", str(DEFAULT_XLSX)))
OUT_DIR = ROOT / "logs"
DOCX_DIR = Path(os.getenv("AI_PATHS_273_DOCX_DIR", str(Path.home() / "Desktop")))
MAX_WORKERS = int(os.getenv("AI_PATHS_273_WORKERS", "20"))
TIMEOUT_SECONDS = int(os.getenv("AI_PATHS_273_TIMEOUT", "180"))
LIMIT = int(os.getenv("AI_PATHS_273_LIMIT", "0"))
USE_FIXED_CUSTOMER_ID = os.getenv("AI_PATHS_273_USE_FIXED_CUSTOMER_ID", "0") == "1"
RETRY_FAILED = int(os.getenv("AI_PATHS_273_RETRY_FAILED", "1"))
INPROCESS_CLIENT = None

BASE_REQUEST = {
    "corp_id": os.getenv("AI_PATHS_TEST_CORP_ID", "ww943af61cd5d2afe4"),
    "user_id": os.getenv("AI_PATHS_TEST_USER_ID", "7294"),
    "wechat": os.getenv("AI_PATHS_TEST_WECHAT", "CS001"),
    "category_id": os.getenv("AI_PATHS_TEST_CATEGORY_ID", "居家产品"),
    "customer_id": os.getenv("AI_PATHS_TEST_CUSTOMER_ID", "20615704"),
    "external_userid": os.getenv("AI_PATHS_TEST_EXTERNAL_USERID", "wmanzqsqaaygjwicitvmos657x39lqtg"),
}

IMAGE_PLACEHOLDERS = {
    "（发送脸部斑点照片）": "我发一张脸部斑点照片，你帮我看看适合什么方向",
    "（发送其他家机构报价截图）": "我发你一张别家机构报价截图，你帮我看看有什么区别",
    "（发送其他家祛斑报价截图）": "我发你一张别家祛斑报价截图，你帮我对比一下",
    "（发送竞品截图）": "我发你一张别家活动截图，你帮我看看",
    "（发送体检报告或病历）": "我发了体检报告，想问问这种情况能不能做",
    "（发送表情包，如大拇指/玫瑰花）": "挺好的",
    "（连续回答城市+困扰+年龄+预算+项目偏好完整度达80%）": (
        "我在上海，脸上老年斑比较多，今年58岁，预算别太高，想先了解淡斑方向"
    ),
}

STORE_KEYWORDS = (
    "门店",
    "地址",
    "位置",
    "附近",
    "导航",
    "停车",
    "营业",
    "哪家",
    "哪里",
    "在哪",
    "城市",
    "机场",
    "厦门",
    "上海",
    "北京",
    "广州",
    "深圳",
    "杭州",
    "成都",
    "重庆",
    "西安",
    "西藏",
)

CITY_NAMES = (
    "北京",
    "上海",
    "广州",
    "深圳",
    "杭州",
    "成都",
    "重庆",
    "西安",
    "厦门",
    "南京",
    "武汉",
    "苏州",
    "天津",
    "郑州",
    "长沙",
    "佛山",
    "东莞",
    "昆明",
    "青岛",
    "济南",
    "合肥",
    "福州",
    "南宁",
    "桂林",
    "西藏",
)

STORE_EXCLUDE_PHRASES = (
    "门店的人",
    "门店怎么有其他客户",
    "一到门店",
    "营业执照",
    "最近有活动",
    "最近身体",
    "最近天气",
    "身体不太好",
)

CITY_ONLY_NEED_LOCATION_HINTS = ("我在", "人在", "我这边在", "这边是", "客户在")

TRANSPORT_FORBIDDEN_TERMS = (
    "车费报销",
    "报销车费",
    "打车报销",
    "打车发票",
    "实报实销",
    "车费补贴",
    "包接送",
    "免费接送",
    "安排接送",
)

FORBIDDEN_TERMS = (
    "焕新体验季",
    "新客专属活动",
    "老带新专属活动",
    "内部活动",
    "大型活动",
    "公司统一通知价",
    "包接送",
    "车费报销",
    "100%见效",
    "根治",
    "绝对安全",
    "保证效果",
    "包效果",
    "一次一定好",
    "AI",
    "机器人",
    "知识库",
    "intent",
    "subflow",
    "reply_brief",
)


def norm(value: Any) -> str:
    return str(value or "").strip()


def load_cases() -> list[dict[str, Any]]:
    wb = openpyxl.load_workbook(INPUT_XLSX, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise RuntimeError(f"Empty workbook: {INPUT_XLSX}")

    header_idx = 0
    headers = [norm(value) for value in rows[header_idx]]
    index = {name: i for i, name in enumerate(headers)}
    required = ["客户阶段", "场景类型", "用户问题"]
    for name in required:
        if name not in index:
            raise RuntimeError(f"Workbook missing column: {name}; headers={headers}")

    logic_col = next((i for i, name in enumerate(headers) if "业务应答" in name), -1)
    sales_col = next((i for i, name in enumerate(headers) if "销冠话术" in name), -1)
    current_stage = ""
    current_scene = ""
    cases: list[dict[str, Any]] = []
    for row_no, row in enumerate(rows[header_idx + 1 :], start=header_idx + 2):
        stage = norm(row[index["客户阶段"]] if index["客户阶段"] < len(row) else "")
        scene = norm(row[index["场景类型"]] if index["场景类型"] < len(row) else "")
        question = norm(row[index["用户问题"]] if index["用户问题"] < len(row) else "")
        if stage:
            current_stage = stage
        if scene:
            current_scene = scene
        if not question:
            continue
        cases.append(
            {
                "index": len(cases) + 1,
                "source_row": row_no,
                "customer_stage": current_stage,
                "scene_type": current_scene,
                "question": question,
                "sent_question": IMAGE_PLACEHOLDERS.get(question, question),
                "business_logic": norm(row[logic_col]) if 0 <= logic_col < len(row) else "",
                "sales_script": norm(row[sales_col]) if 0 <= sales_col < len(row) else "",
            }
        )
        if LIMIT and len(cases) >= LIMIT:
            break
    return cases


def build_payload(case: dict[str, Any], run_stamp: str) -> dict[str, Any]:
    if USE_FIXED_CUSTOMER_ID:
        customer_id = BASE_REQUEST["customer_id"]
        external_userid = BASE_REQUEST["external_userid"]
    else:
        customer_id = f"codex273_{run_stamp}_{case['index']}_{uuid.uuid4().hex[:6]}"
        external_userid = customer_id
    return {
        "corp_id": BASE_REQUEST["corp_id"],
        "user_id": BASE_REQUEST["user_id"],
        "wechat": BASE_REQUEST["wechat"],
        "category_id": BASE_REQUEST["category_id"],
        "customer_id": customer_id,
        "external_userid": external_userid,
        "content": case["sent_question"],
        "conversation_history": [],
    }


def call_case_once(case: dict[str, Any], run_stamp: str) -> dict[str, Any]:
    payload = build_payload(case, run_stamp)
    started = time.perf_counter()
    status = 0
    error = ""
    data: dict[str, Any] = {}
    if INPROCESS:
        try:
            client = get_inprocess_client()
            response = client.post("/reply", json=payload, timeout=TIMEOUT_SECONDS)
            status = response.status_code
            try:
                data = response.json()
            except Exception:
                data = {"raw_error": response.text}
            if status >= 400:
                error = f"HTTP {status}"
        except Exception as exc:
            error = type(exc).__name__
            data = {"raw_error": str(exc)}
    else:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        request = urllib.request.Request(
            API_URL,
            data=body,
            headers={"Content-Type": "application/json; charset=utf-8"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(request, timeout=TIMEOUT_SECONDS) as response:
                status = response.status
                data = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            status = exc.code
            raw = exc.read().decode("utf-8", errors="replace")
            try:
                data = json.loads(raw)
            except Exception:
                data = {"raw_error": raw}
            error = f"HTTP {exc.code}"
        except Exception as exc:
            error = type(exc).__name__
            data = {"raw_error": str(exc)}
    elapsed_ms = int((time.perf_counter() - started) * 1000)
    replies, reply_types, handoff_reason = extract_reply_messages(data)
    meta = data.get("meta") if isinstance(data.get("meta"), dict) else {}
    tool_result_keys = [norm(item) for item in meta.get("tool_result_keys") or []]
    tool_call_names = extract_tool_call_names(meta)
    store_required = is_store_required(case)
    store_called = any("store" in item.lower() or "门店" in item for item in [*tool_result_keys, *tool_call_names])
    forbidden_hits = [term for term in FORBIDDEN_TERMS if term in " ".join(replies)]
    judgement = judge(
        status=status,
        error=error,
        replies=replies,
        forbidden_hits=forbidden_hits,
        store_required=store_required,
        store_called=store_called,
        case=case,
    )
    return {
        **case,
        "status": status,
        "error": error,
        "elapsed_ms": elapsed_ms,
        "reply_1": replies[0] if replies else "",
        "reply_2": replies[1] if len(replies) > 1 else "",
        "reply_count": len(replies),
        "reply_types": reply_types,
        "handoff_reason": handoff_reason,
        "log_id": norm(data.get("request_id") or data.get("trace_id") or data.get("execute_id")),
        "scene": norm(data.get("scene")),
        "intent": norm(data.get("intent")),
        "subflow": norm(data.get("subflow")),
        "policy_family_id": norm(meta.get("policy_family_id")),
        "exact_policy_id": norm(meta.get("exact_policy_id")),
        "active_scene_id": norm(meta.get("active_scene_id")),
        "canonical_sales_reply": norm(meta.get("canonical_sales_reply")),
        "sales_script_similarity": meta.get("sales_script_similarity", ""),
        "tool_result_keys": tool_result_keys,
        "tool_call_names": tool_call_names,
        "store_required": store_required,
        "store_called": store_called,
        "forbidden_hits": forbidden_hits,
        "judgement": judgement,
        "token_usage": meta.get("token_usage") if isinstance(meta.get("token_usage"), dict) else {},
    }


def get_inprocess_client():
    global INPROCESS_CLIENT
    if INPROCESS_CLIENT is None:
        from fastapi.testclient import TestClient

        from app.main import app

        client = TestClient(app)
        client.__enter__()
        INPROCESS_CLIENT = client
    return INPROCESS_CLIENT


def close_inprocess_client() -> None:
    global INPROCESS_CLIENT
    if INPROCESS_CLIENT is not None:
        INPROCESS_CLIENT.__exit__(None, None, None)
        INPROCESS_CLIENT = None


def should_retry(result: dict[str, Any]) -> bool:
    if int(result.get("status") or 0) != 200:
        return True
    if result.get("error"):
        return True
    if not norm(result.get("reply_1")) and not norm(result.get("handoff_reason")):
        return True
    return False


def call_case(case: dict[str, Any], run_stamp: str) -> dict[str, Any]:
    last_result: dict[str, Any] | None = None
    attempts = max(1, RETRY_FAILED + 1)
    for attempt in range(attempts):
        result = call_case_once(case, run_stamp)
        result["attempts"] = attempt + 1
        last_result = result
        if not should_retry(result):
            return result
        if attempt + 1 < attempts:
            time.sleep(min(3, 0.6 * (attempt + 1)))
    return last_result or call_case_once(case, run_stamp)


def extract_reply_messages(data: dict[str, Any]) -> tuple[list[str], list[str], str]:
    replies: list[str] = []
    reply_types: list[str] = []
    handoff_reason = ""
    for item in data.get("reply_messages") or []:
        if not isinstance(item, dict):
            continue
        msg_type = norm(item.get("type"))
        reply_types.append(msg_type)
        content = item.get("content")
        if msg_type == "text":
            if isinstance(content, dict):
                text = norm(content.get("text"))
            else:
                text = norm(content)
            if text:
                replies.append(text)
        elif msg_type == "human_handoff":
            if isinstance(content, dict):
                handoff_reason = norm(content.get("handoff_reason"))
            else:
                handoff_reason = norm(content)
    return replies, reply_types, handoff_reason


def extract_tool_call_names(meta: dict[str, Any]) -> list[str]:
    names: list[str] = []
    for call in meta.get("tool_calls") or []:
        if not isinstance(call, dict):
            continue
        name = norm(call.get("name"))
        if name:
            names.append(name)
        nested = call.get("nested_calls") or []
        if isinstance(nested, list):
            for item in nested:
                if isinstance(item, dict) and norm(item.get("name")):
                    names.append(norm(item.get("name")))
    return names


def is_store_required(case: dict[str, Any]) -> bool:
    question = norm(case.get("question"))
    scene_type = norm(case.get("scene_type"))
    customer_stage = norm(case.get("customer_stage"))
    text = f"{customer_stage} {scene_type} {question}"
    if any(phrase in question for phrase in STORE_EXCLUDE_PHRASES):
        return False
    if "主动暴露画像-城市" in scene_type:
        return True
    if any(city in question for city in CITY_NAMES) and (
        any(hint in question for hint in CITY_ONLY_NEED_LOCATION_HINTS)
        or "有店" in question
        or "门店" in question
        or "附近" in question
        or "地址" in question
    ):
        return True
    return any(keyword in text for keyword in STORE_KEYWORDS)


def is_city_only_profile(case: dict[str, Any]) -> bool:
    question = norm(case.get("question"))
    scene_type = norm(case.get("scene_type"))
    if "主动暴露画像-城市" in scene_type:
        return True
    return any(city in question for city in CITY_NAMES) and any(hint in question for hint in CITY_ONLY_NEED_LOCATION_HINTS)


def is_transport_support_question(case: dict[str, Any]) -> bool:
    question = norm(case.get("question"))
    return any(term in question for term in ("车费", "接送", "派车", "路费", "打车"))


def judge(
    *,
    status: int,
    error: str,
    replies: list[str],
    forbidden_hits: list[str],
    store_required: bool,
    store_called: bool,
    case: dict[str, Any] | None = None,
) -> str:
    issues: list[str] = []
    joined_reply = " ".join(replies)
    if status != 200 or error:
        issues.append(f"接口异常：{error or status}")
    if not replies:
        issues.append("无客户可见回复")
    if forbidden_hits:
        issues.append("命中禁用词：" + "、".join(forbidden_hits[:5]))
    if store_required and not store_called:
        issues.append("门店类未确认调用 store_lookup")
    if case and is_city_only_profile(case) and replies:
        if not any(term in joined_reply for term in ("哪个区", "哪一区", "附近什么地标", "附近地标", "哪个位置", "哪个区域", "什么地标")):
            issues.append("城市画像未继续追问区/地标")
        if any(term in joined_reply for term in ("地址", "营业时间", "最近门店")):
            issues.append("城市画像过早输出具体门店事实")
    if case and is_transport_support_question(case) and replies:
        transport_hits = [term for term in TRANSPORT_FORBIDDEN_TERMS if term in joined_reply]
        if transport_hits:
            issues.append("交通服务回复含风险词：" + "、".join(transport_hits[:5]))
    if replies and len(replies[0]) > 120:
        issues.append("首条偏长")
    if len(replies) > 2:
        issues.append("回复条数偏多")
    return "通过" if not issues else "；".join(issues)


def run_all(cases: list[dict[str, Any]], run_stamp: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    mode = "inprocess" if INPROCESS else "online"
    print(
        f"start {mode} 273 cases={len(cases)} workers={MAX_WORKERS} timeout={TIMEOUT_SECONDS}s api={API_URL}",
        flush=True,
    )
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_map = {executor.submit(call_case, case, run_stamp): case for case in cases}
        completed = 0
        for future in as_completed(future_map):
            completed += 1
            try:
                result = future.result()
            except Exception as exc:
                case = future_map[future]
                result = {
                    **case,
                    "status": 0,
                    "error": type(exc).__name__,
                    "elapsed_ms": 0,
                    "reply_1": "",
                    "reply_2": "",
                    "reply_count": 0,
                    "reply_types": [],
                    "handoff_reason": "",
                    "log_id": "",
                    "policy_family_id": "",
                    "exact_policy_id": "",
                    "active_scene_id": "",
                    "tool_result_keys": [],
                    "tool_call_names": [],
                    "store_required": is_store_required(case),
                    "store_called": False,
                    "forbidden_hits": [],
                    "judgement": f"执行异常：{type(exc).__name__}",
                    "token_usage": {},
                }
            results.append(result)
            if completed % 20 == 0 or completed == len(cases):
                print(f"completed {completed}/{len(cases)}", flush=True)
    return sorted(results, key=lambda item: int(item.get("index") or 0))


def write_jsonl(results: list[dict[str, Any]], run_stamp: str) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / f"ai_customer_reply_273_online_results_{run_stamp}.jsonl"
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        for item in results:
            handle.write(json.dumps(item, ensure_ascii=False) + "\n")
    return path


def write_markdown(results: list[dict[str, Any]], run_stamp: str, summary: dict[str, Any]) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / f"ai_customer_reply_273_online_report_{run_stamp}.md"
    lines = [
        "# AI客服273条线上全量测试报告",
        "",
        f"- 运行时间：{summary['generated_at']}",
        f"- 模式：`{'本地进程' if INPROCESS else '线上接口'}`",
        f"- 接口：`{API_URL if not INPROCESS else '/reply (TestClient)'}`",
        f"- 输入表：`{INPUT_XLSX}`",
        f"- 总数：{summary['total']}",
        f"- 通过：{summary['passed']}",
        f"- 需优化/异常：{summary['failed']}",
        f"- 门店类样本：{summary['store_required']}",
        f"- 门店类已调用工具：{summary['store_called']}",
        f"- 门店类缺工具：{summary['store_missing']}",
        f"- 超时/接口异常：{summary['http_or_timeout_errors']}",
        f"- 平均耗时：{summary['avg_seconds']}s",
        f"- 总耗时：{summary['total_seconds']}s",
        "",
        "| 客户阶段 | 场景类型 | 用户问题 | AI实际回复（第1条） | AI引导回复（第2条） | 日志id | 工具调用 | 门店工具 | 评判 |",
        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
    ]
    for item in results:
        lines.append(
            "| "
            + " | ".join(
                [
                    md_cell(item.get("customer_stage")),
                    md_cell(item.get("scene_type")),
                    md_cell(item.get("question")),
                    md_cell(item.get("reply_1")),
                    md_cell(item.get("reply_2")),
                    md_cell(item.get("log_id")),
                    md_cell(", ".join(item.get("tool_result_keys") or item.get("tool_call_names") or [])),
                    md_cell(store_status(item)),
                    md_cell(item.get("judgement")),
                ]
            )
            + " |"
        )
    path.write_text("\n".join(lines), encoding="utf-8-sig", newline="\n")
    return path


def md_cell(value: Any) -> str:
    return norm(value).replace("\\", "\\\\").replace("|", "/").replace("\r", "").replace("\n", "<br>")


def store_status(item: dict[str, Any]) -> str:
    if not item.get("store_required"):
        return "不要求"
    return "已调用" if item.get("store_called") else "缺失"


def summarize(results: list[dict[str, Any]], total_seconds: float) -> dict[str, Any]:
    passed = sum(1 for item in results if item.get("judgement") == "通过")
    store_required = [item for item in results if item.get("store_required")]
    elapsed = [int(item.get("elapsed_ms") or 0) for item in results if item.get("elapsed_ms")]
    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total": len(results),
        "passed": passed,
        "failed": len(results) - passed,
        "store_required": len(store_required),
        "store_called": sum(1 for item in store_required if item.get("store_called")),
        "store_missing": sum(1 for item in store_required if not item.get("store_called")),
        "http_or_timeout_errors": sum(1 for item in results if item.get("status") != 200 or item.get("error")),
        "avg_seconds": round(mean(elapsed) / 1000, 1) if elapsed else 0,
        "total_seconds": round(total_seconds, 1),
    }


def set_cell_text(cell, text: Any, *, size: int = 8, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(norm(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_docx(results: list[dict[str, Any]], run_stamp: str, summary: dict[str, Any]) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    path = DOCX_DIR / f"AI客服273条线上全量测试报告_{run_stamp}.docx"
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title_run = title.add_run("AI客服273条线上全量测试报告")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    overview = doc.add_paragraph()
    overview.add_run(
        f"运行时间：{summary['generated_at']}；模式：{'本地进程' if INPROCESS else '线上接口'}；"
        f"接口：{API_URL if not INPROCESS else '/reply (TestClient)'}；输入表：{INPUT_XLSX}"
    )

    summary_table = doc.add_table(rows=2, cols=8)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = "Table Grid"
    headers = ["总数", "通过", "需优化/异常", "门店类", "门店已调用", "门店缺工具", "接口/超时异常", "平均耗时"]
    values = [
        summary["total"],
        summary["passed"],
        summary["failed"],
        summary["store_required"],
        summary["store_called"],
        summary["store_missing"],
        summary["http_or_timeout_errors"],
        f"{summary['avg_seconds']}s",
    ]
    for i, header in enumerate(headers):
        set_cell_text(summary_table.cell(0, i), header, size=9, bold=True)
        shade_cell(summary_table.cell(0, i), "D9EAF7")
        set_cell_text(summary_table.cell(1, i), values[i], size=9)

    doc.add_paragraph(
        "说明：门店类判定基于客户阶段、场景类型、用户问题中的门店/地址/位置/附近/导航/营业/停车/城市等关键词；"
        "工具调用依据接口 meta.tool_result_keys 和 meta.tool_calls 统计。"
    )

    detail_headers = [
        "序号",
        "客户阶段",
        "场景类型",
        "用户问题",
        "AI实际回复（第1条）",
        "AI引导回复（第2条）",
        "日志id",
        "工具调用",
        "门店工具",
        "评判",
    ]
    table = doc.add_table(rows=1, cols=len(detail_headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(detail_headers):
        set_cell_text(table.cell(0, i), header, size=8, bold=True)
        shade_cell(table.cell(0, i), "1F4E79")
        for run in table.cell(0, i).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for item in results:
        row = table.add_row().cells
        values = [
            item.get("index"),
            item.get("customer_stage"),
            item.get("scene_type"),
            item.get("question"),
            item.get("reply_1"),
            item.get("reply_2"),
            item.get("log_id"),
            ", ".join(item.get("tool_result_keys") or item.get("tool_call_names") or []),
            store_status(item),
            item.get("judgement"),
        ]
        for i, value in enumerate(values):
            set_cell_text(row[i], value, size=7)
        if item.get("judgement") != "通过":
            shade_cell(row[-1], "FCE4D6")
        elif item.get("store_required"):
            shade_cell(row[-2], "E2F0D9")

    doc.add_page_break()
    doc.add_heading("主要问题样本", level=1)
    failed = [item for item in results if item.get("judgement") != "通过"][:40]
    if not failed:
        doc.add_paragraph("本轮基础评判未发现异常样本。")
    else:
        for item in failed:
            p = doc.add_paragraph()
            p.add_run(f"{item.get('index')}. {item.get('question')}").bold = True
            doc.add_paragraph(f"评判：{item.get('judgement')}")
            doc.add_paragraph(f"回复：{item.get('reply_1')} {item.get('reply_2')}".strip())
            doc.add_paragraph(f"日志ID：{item.get('log_id')}")

    doc.save(path)
    verify_docx(path)
    return path


def verify_docx(path: Path) -> None:
    loaded = Document(path)
    if not loaded.paragraphs:
        raise RuntimeError(f"DOCX verification failed: no paragraphs in {path}")
    if not loaded.tables:
        raise RuntimeError(f"DOCX verification failed: no tables in {path}")


def main() -> None:
    started = time.perf_counter()
    run_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    reuse_jsonl = os.getenv("AI_PATHS_273_REUSE_JSONL", "").strip()
    try:
        if reuse_jsonl:
            results = load_existing_results(Path(reuse_jsonl))
        else:
            cases = load_cases()
            results = run_all(cases, run_stamp)
    finally:
        close_inprocess_client()
    total_seconds = float(os.getenv("AI_PATHS_273_REUSE_TOTAL_SECONDS", "0") or 0) or (time.perf_counter() - started)
    summary = summarize(results, total_seconds)
    jsonl_path = write_jsonl(results, run_stamp)
    md_path = write_markdown(results, run_stamp, summary)
    docx_path = write_docx(results, run_stamp, summary)
    print("SUMMARY=" + json.dumps(summary, ensure_ascii=False))
    print("JSONL=" + str(jsonl_path.resolve()))
    print("MD=" + str(md_path.resolve()))
    print("DOCX=" + str(docx_path.resolve()))


def load_existing_results(path: Path) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        item = json.loads(line)
        item["store_required"] = is_store_required(item)
        item["store_called"] = any(
            "store" in norm(value).lower() or "门店" in norm(value)
            for value in [*(item.get("tool_result_keys") or []), *(item.get("tool_call_names") or [])]
        )
        forbidden_hits = [term for term in FORBIDDEN_TERMS if term in f"{item.get('reply_1','')} {item.get('reply_2','')}"]
        item["forbidden_hits"] = forbidden_hits
        item["judgement"] = judge(
            status=int(item.get("status") or 0),
            error=norm(item.get("error")),
            replies=[text for text in [norm(item.get("reply_1")), norm(item.get("reply_2"))] if text],
            forbidden_hits=forbidden_hits,
            store_required=bool(item.get("store_required")),
            store_called=bool(item.get("store_called")),
            case=item,
        )
        results.append(item)
    return sorted(results, key=lambda row: int(row.get("index") or 0))


if __name__ == "__main__":
    main()
